import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from get_week_dates import get_week_dates
from create_reports_directory import create_reports_directory

def create_word_document(report, personal_informations):
    name, ausbildungsberuf, geschaeftsbereich, ausbildungsjahr = personal_informations

    # Get the reports directory
    reports_folder = create_reports_directory()
    
    # Create new document with A4 size
    doc = Document()
    
    # Add header table for the first section
    header_table = doc.add_table(rows=2, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Set column widths
    header_table.columns[0].width = Cm(10)
    header_table.columns[1].width = Cm(10)
    
    # First row
    cell1 = header_table.cell(0, 0)
    p1 = cell1.paragraphs[0]
    p1.add_run('Ausbildungsnachweis Nr. ').bold = True
    p1.add_run('46')
    
    cell2 = header_table.cell(0, 1)
    p2 = cell2.paragraphs[0]
    p2.add_run(name).bold = True
    p2.add_run('\n(Name der Auszubildenden)')
    
    # Second row
    start_date, end_date = get_week_dates()
    cell3 = header_table.cell(1, 0)
    p3 = cell3.paragraphs[0]
    p3.add_run(f'vom {start_date} bis {end_date}')
    
    cell4 = header_table.cell(1, 1)
    p4 = cell4.paragraphs[0]
    p4.add_run(ausbildungsberuf).bold = True
    p4.add_run('\n(Ausbildungsberuf)')
    
    # Add name and year
    doc.add_paragraph()
    name_para = doc.add_paragraph()
    name_para.add_run(f"{name} IT").bold = True
    name_para.add_run('\n(Geschäftsbereich)')
    
    year_para = doc.add_paragraph()
    year_para.add_run(ausbildungsjahr).bold = True
    year_para.add_run('\n(Ausbildungsjahr)')
    
    # Add main heading
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run('AUSBILDUNGSINHALTE').bold = True
    
    # Create activity table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths (adjusted to match template)
    table.columns[0].width = Cm(2.5)     # Tag column - narrower
    table.columns[1].width = Cm(15)      # Activities column - much wider
    table.columns[2].width = Cm(1.5)     # Std column - very narrow
    
    # Set table headers
    headers = table.rows[0].cells
    headers[0].text = "Tag"
    headers[1].text = "Ausgeführte Arbeiten, Unterricht usw."
    headers[2].text = "Std"
    
    # Fill table with activities and hours
    days = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag"]
    for i, day in enumerate(days):
        row = table.rows[i+1]
        row.cells[0].text = day
        # Remove "Am" from the beginning since it's in a separate column now
        activity_text = report.get(day, "").replace(f"Am {day} : ", "")
        row.cells[1].text = activity_text
        row.cells[2].text = "8"  # Standard hours
        
        # Set vertical alignment for all cells in the row
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # Add signature section
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Create signature table
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.style = 'Table Grid'
    
    # Add signature lines
    sigs = ['Datum\nAuszubildende', 'Datum\nAusbilder/in', 'Datum\ngesetzl. Vertreter']
    for i, sig in enumerate(sigs):
        sig_table.cell(0, i).text = sig
    
    # Save document
    filename = f'Ausbildungsnachweis_{start_date}_bis_{end_date}.docx'
    filepath = os.path.join(reports_folder, filename)
    doc.save(filepath)
    
    print(f'Dokument erfolgreich erstellt: {filepath}')