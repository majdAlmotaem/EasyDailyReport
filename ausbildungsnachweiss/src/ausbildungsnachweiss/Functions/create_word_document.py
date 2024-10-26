import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from get_week_dates import get_week_dates
from create_reports_directory import create_reports_directory

def create_word_document(report, personal_informations):
    name, ausbildungsberuf, geschaeftsbereich, ausbildungsjahr = personal_informations

    # Get the reports directory
    reports_folder = create_reports_directory()
    
    # Create new document
    doc = Document()
    
    # Title
    heading = doc.add_heading('Ausbildungsnachweis', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_format = heading.runs[0].font
    heading_format.size = Pt(24)
    heading_format.bold = True
    doc.add_paragraph("")
    
    # Information Section with filled data
    doc.add_paragraph(f"Name der Auszubildenden: {name}")
    doc.add_paragraph(f"Ausbildungsberuf: {ausbildungsberuf}")
    doc.add_paragraph(f"Geschäftsbereich: {geschaeftsbereich}")
    doc.add_paragraph(f"Ausbildungsjahr: {ausbildungsjahr}")
    
    # Add dates
    start_date, end_date = get_week_dates()
    doc.add_paragraph(f"Zeitraum vom: {start_date} bis: {end_date}")
    
    # Empty line before table
    doc.add_paragraph("")
    
    # Create activity table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(5.5)
    
    # Set table headers
    table.cell(0, 0).text = "Tag"
    table.cell(0, 1).text = "Ausgeführte Arbeiten, Unterricht usw."
    
    # Fill table with days and activities
    days = ["Montag", "Dienstag", "Mittwoch", "Donnerstag", "Freitag"]
    for i, day in enumerate(days, 1):
        table.cell(i, 0).text = day
        table.cell(i, 1).text = report.get(day, "")
    
    # Empty line before signatures
    doc.add_paragraph("")
    
    # Signature Section
    doc.add_paragraph("Datum: __________________ Auszubildende/r: __________________")
    doc.add_paragraph("Datum: __________________ Ausbilder/in: __________________")
    doc.add_paragraph("Datum: __________________ gesetzl. Vertreter: __________________")
    
    # Save document
    filename = f'Ausbildungsnachweis_{start_date}_bis_{end_date}.docx'
    filepath = os.path.join(reports_folder, filename)
    doc.save(filepath)
    
    print(f'Dokument erfolgreich erstellt: {filepath}')